import io
import logging

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import get_current_user
from app.db.base import get_db
from app.models.governance import ComplianceFinding, GovernanceDocument
from app.models.user import User
from app.schemas.governance import (
    ComplianceFindingList,
    ComplianceFindingResponse,
    GovernanceAnalyzeRequest,
    GovernanceDocumentList,
    GovernanceDocumentResponse,
)
from app.services.sqs_publisher import SQSPublisher

logger = logging.getLogger(__name__)

router = APIRouter()

_publisher = SQSPublisher()

def _extract_text(filename: str, content: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".docx"):
        from docx import Document

        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    if lower.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    return content.decode("utf-8", errors="replace")

@router.post("/{org_login}/governance/documents", response_model=GovernanceDocumentResponse)
async def upload_governance_document(
    org_login: str,
    doc_type: str,
    title: str,
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> GovernanceDocumentResponse:
    if doc_type not in ("governance_policy", "app_profile"):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="doc_type must be governance_policy or app_profile")

    content = await file.read()
    try:
        raw_text = _extract_text(file.filename or "", content)
    except Exception as exc:
        logger.warning("Failed to extract text from %s: %s", file.filename, exc)
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Could not parse document")

    document = GovernanceDocument(
        org_login=org_login,
        doc_type=doc_type,
        title=title,
        source_filename=file.filename,
        raw_text=raw_text,
        uploaded_by=user.id,
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)

    await _publisher.publish({
        "event_type": "extract_governance_requirements",
        "document_id": str(document.id),
    })

    return GovernanceDocumentResponse.model_validate(document)

@router.get("/{org_login}/governance/documents", response_model=GovernanceDocumentList)
async def list_governance_documents(
    org_login: str,
    _user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> GovernanceDocumentList:
    result = await db.execute(
        select(GovernanceDocument)
        .where(GovernanceDocument.org_login == org_login)
        .order_by(GovernanceDocument.created_at.desc())
    )
    documents = result.scalars().all()
    return GovernanceDocumentList(documents=[GovernanceDocumentResponse.model_validate(d) for d in documents])

@router.post("/{org_login}/repos/{repo_name}/governance/analyze")
async def analyze_governance(
    org_login: str,
    repo_name: str,
    body: GovernanceAnalyzeRequest,
    _user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> dict:
    if body.mode not in ("framework", "document"):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="mode must be framework or document")
    if body.mode == "framework" and not body.framework:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="framework is required for mode=framework")
    if body.mode == "document" and not body.document_id:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="document_id is required for mode=document")

    await _publisher.publish({
        "event_type": "run_governance_analysis",
        "org_login": org_login,
        "repo_name": repo_name,
        "ref": body.ref,
        "mode": body.mode,
        "framework": body.framework,
        "document_id": str(body.document_id) if body.document_id else None,
    })
    return {"status": "enqueued", "org_login": org_login, "repo_name": repo_name, "mode": body.mode}

@router.get("/{org_login}/repos/{repo_name}/governance/findings", response_model=ComplianceFindingList)
async def get_compliance_findings(
    org_login: str,
    repo_name: str,
    _user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> ComplianceFindingList:
    result = await db.execute(
        select(ComplianceFinding)
        .where(ComplianceFinding.org_login == org_login, ComplianceFinding.repo_name == repo_name)
        .order_by(ComplianceFinding.computed_at.desc())
    )
    findings = result.scalars().all()
    return ComplianceFindingList(findings=[ComplianceFindingResponse.model_validate(f) for f in findings])
